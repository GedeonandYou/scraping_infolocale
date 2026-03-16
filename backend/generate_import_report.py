"""Génère un rapport Word pour l'import des données CSV Infolocale dans GEDEON Admin."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date


def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule."""
    from lxml import etree
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")


def add_styled_table(doc, headers, rows, col_widths=None):
    """Crée un tableau stylisé."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E74B5")

    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generate_report():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ==========================================================
    # PAGE DE TITRE
    # ==========================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GUIDE D'IMPORT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Import des événements Infolocale\n"
        "dans GEDEON Admin"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Fichier source : infolocale_20260315_8.csv\n"
        f"Date : {date.today().strftime('%d/%m/%Y')}\n"
        f"Destination : GEDEON Admin (admin-pwa)"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()

    # ==========================================================
    # TABLE DES MATIÈRES
    # ==========================================================
    doc.add_heading("Table des matières", level=1)
    toc_items = [
        "1. Vue d'ensemble",
        "2. Analyse du fichier CSV source",
        "3. Champs attendus par GEDEON Admin",
        "4. Correspondance des champs (mapping)",
        "5. Transformations de données nécessaires",
        "6. Procédure d'import pas à pas",
        "7. Utilisation du mapping IA",
        "8. Champs non importables",
        "9. Points d'attention",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ==========================================================
    # 1. VUE D'ENSEMBLE
    # ==========================================================
    doc.add_heading("1. Vue d'ensemble", level=1)

    doc.add_paragraph(
        "Ce document décrit la procédure pour importer les événements extraits "
        "d'Infolocale (via l'API Algolia) dans l'application GEDEON Admin. "
        "Le fichier CSV généré par le script export_events.py contient 5 000 "
        "événements avec 22 colonnes de données."
    )

    doc.add_heading("1.1 Fichier source", level=2)
    add_styled_table(
        doc,
        ["Propriété", "Valeur"],
        [
            ["Nom du fichier", "infolocale_20260315_8.csv"],
            ["Emplacement", "backend/data/exports/"],
            ["Nombre de lignes", "5 000 événements"],
            ["Nombre de colonnes", "22"],
            ["Encodage", "UTF-8"],
            ["Séparateur", "Virgule (,)"],
            ["Format des dates", "JJ/MM/AAAA HH:MM:SS"],
        ],
        col_widths=[5, 9],
    )

    doc.add_paragraph()

    doc.add_heading("1.2 Destination", level=2)
    add_styled_table(
        doc,
        ["Propriété", "Valeur"],
        [
            ["Application", "GEDEON Admin (admin-pwa)"],
            ["URL", "https://gedeon-admin.onrender.com"],
            ["Méthode d'import", "Interface web (3 étapes)"],
            ["Formats acceptés", "CSV, Excel (.xlsx/.xls), JSON"],
            ["API backend", "POST /api/admin/import-full (SSE streaming)"],
            ["Mapping", "Manuel ou assisté par IA"],
        ],
        col_widths=[5, 9],
    )

    doc.add_page_break()

    # ==========================================================
    # 2. ANALYSE DU CSV SOURCE
    # ==========================================================
    doc.add_heading("2. Analyse du fichier CSV source", level=1)

    doc.add_paragraph(
        "Le fichier CSV exporté depuis Algolia contient les 22 colonnes suivantes :"
    )

    add_styled_table(
        doc,
        ["#", "Colonne CSV", "Type", "Exemple"],
        [
            ["1", "objectID", "UUID", "e036a342-dffa-3e2a-87e1-..."],
            ["2", "titre", "Texte", "Swin'Gum au Relax"],
            ["3", "texte", "Texte long", "Dans la famille Gîte et Cover..."],
            ["4", "categorie", "Texte", "Concert, spectacle musical"],
            ["5", "genre", "Texte", "Chanson"],
            ["6", "rubrique_lvl0", "Texte", "Concerts, spectacles"],
            ["7", "rubrique_lvl1", "Texte", "Concerts, spectacles > Concert..."],
            ["8", "date_debut", "Date+Heure", "15/03/2026 20:00:00"],
            ["9", "date_fin", "Date+Heure", "15/03/2026 23:00:00"],
            ["10", "lieu_nom", "Texte", "Lille"],
            ["11", "lieu_adresse", "Texte", "Place de la Nouvelle Aventure, Lille"],
            ["12", "latitude", "Float", "50.63171768188477"],
            ["13", "longitude", "Float", "3.047832727432251"],
            ["14", "photo_url", "URL", "https://photos.infolocale.fr/..."],
            ["15", "url", "URL", "https://www.infolocale.fr/..."],
            ["16", "annule", "Booléen", "False"],
            ["17", "complet", "Booléen", "False"],
            ["18", "permanent", "Booléen", "False"],
            ["19", "source", "Texte", "IL"],
            ["20", "provider", "Texte", "IL / OA"],
            ["21", "tarifs", "Texte", "0.0"],
            ["22", "ages", "Texte", "Tout Public"],
            ["23", "accessibilites", "Texte (|)", "Visuel|Mental|Moteur"],
        ],
        col_widths=[1, 3.5, 3, 6.5],
    )

    doc.add_page_break()

    # ==========================================================
    # 3. CHAMPS GEDEON ADMIN
    # ==========================================================
    doc.add_heading("3. Champs attendus par GEDEON Admin", level=1)

    doc.add_paragraph(
        "L'import GEDEON Admin attend les champs suivants. Les 3 premiers sont "
        "obligatoires, les autres sont optionnels."
    )

    doc.add_heading("3.1 Champs obligatoires", level=2)
    add_styled_table(
        doc,
        ["Champ BD", "Type", "Description"],
        [
            ["title", "Texte", "Titre de l'événement"],
            ["begin_date", "Date", "Date de début (AAAA-MM-JJ ou JJ-MM-AAAA)"],
            ["location_name", "Texte", "Nom du lieu"],
        ],
        col_widths=[4, 3, 7],
    )

    doc.add_paragraph()

    doc.add_heading("3.2 Champs optionnels", level=2)
    add_styled_table(
        doc,
        ["Champ BD", "Type", "Description"],
        [
            ["end_date", "Date", "Date de fin"],
            ["start_time", "Heure", "Heure de début (HH:MM)"],
            ["end_time", "Heure", "Heure de fin (HH:MM)"],
            ["address", "Texte", "Adresse du lieu"],
            ["city", "Texte", "Ville"],
            ["zipcode", "Texte", "Code postal"],
            ["state", "Texte", "Département / Région"],
            ["country", "Texte", "Pays"],
            ["latitude", "Float", "Latitude GPS"],
            ["longitude", "Float", "Longitude GPS"],
            ["category", "Texte", "Catégorie (Concert, Théâtre, Sport...)"],
            ["description", "Texte", "Description longue"],
            ["organizer", "Texte", "Organisateur"],
            ["pricing", "Texte", "Tarification (Gratuit, 10\u20ac-25\u20ac...)"],
            ["website", "URL", "Site web de l'événement"],
            ["tags", "Tableau", "Mots-clés (séparés par virgule ou pipe)"],
            ["artists", "Tableau", "Artistes"],
            ["sponsors", "Tableau", "Sponsors"],
            ["is_private", "Booléen", "Événement privé (true/false)"],
        ],
        col_widths=[3.5, 2.5, 8],
    )

    doc.add_page_break()

    # ==========================================================
    # 4. CORRESPONDANCE DES CHAMPS
    # ==========================================================
    doc.add_heading("4. Correspondance des champs (mapping)", level=1)

    doc.add_paragraph(
        "Le tableau ci-dessous définit la correspondance entre les colonnes du "
        "fichier CSV Infolocale et les champs attendus par GEDEON Admin."
    )

    doc.add_heading("4.1 Mapping recommandé", level=2)
    add_styled_table(
        doc,
        ["Colonne CSV", "\u2192", "Champ GEDEON", "Transformation", "Obligatoire"],
        [
            ["titre", "\u2192", "title", "Aucune", "OUI"],
            ["date_debut", "\u2192", "begin_date", "Extraire la date (JJ/MM/AAAA)", "OUI"],
            ["lieu_nom", "\u2192", "location_name", "Aucune", "OUI"],
            ["date_fin", "\u2192", "end_date", "Extraire la date (JJ/MM/AAAA)", "Non"],
            ["date_debut", "\u2192", "start_time", "Extraire l'heure (HH:MM)", "Non"],
            ["date_fin", "\u2192", "end_time", "Extraire l'heure (HH:MM)", "Non"],
            ["lieu_adresse", "\u2192", "address", "Aucune", "Non"],
            ["latitude", "\u2192", "latitude", "Aucune", "Non"],
            ["longitude", "\u2192", "longitude", "Aucune", "Non"],
            ["categorie", "\u2192", "category", "Aucune", "Non"],
            ["texte", "\u2192", "description", "Aucune", "Non"],
            ["tarifs", "\u2192", "pricing", "Aucune", "Non"],
            ["url", "\u2192", "website", "Aucune", "Non"],
            ["ages", "\u2192", "tags", "Aucune", "Non"],
        ],
        col_widths=[3, 1, 3, 4.5, 2.5],
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Note importante : ")
    run.bold = True
    p.add_run(
        "La colonne date_debut contient à la fois la date ET l'heure "
        "(format : 15/03/2026 20:00:00). Le mapping IA de GEDEON sait "
        "généralement séparer la date et l'heure automatiquement. "
        "En mapping manuel, mapper date_debut vers begin_date suffira "
        "car le backend accepte le format JJ/MM/AAAA."
    )

    doc.add_page_break()

    # ==========================================================
    # 5. TRANSFORMATIONS
    # ==========================================================
    doc.add_heading("5. Transformations de données nécessaires", level=1)

    doc.add_heading("5.1 Dates et heures", level=2)
    doc.add_paragraph(
        "Le CSV contient les dates au format \u00ab JJ/MM/AAAA HH:MM:SS \u00bb "
        "(ex : 15/03/2026 20:00:00). GEDEON Admin attend :"
    )
    for item in [
        "begin_date : format AAAA-MM-JJ ou JJ-MM-AAAA (la partie date)",
        "start_time : format HH:MM (la partie heure)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Le mapping IA détecte automatiquement ce format composite et effectue "
        "la séparation. En import manuel, le backend de GEDEON tente aussi de "
        "parser les formats courants."
    )

    doc.add_heading("5.2 Champs tableau (accessibilites)", level=2)
    doc.add_paragraph(
        "La colonne accessibilites utilise le séparateur pipe (|) : "
        "\u00ab Visuel|Mental|Moteur \u00bb. GEDEON Admin accepte les séparateurs "
        "virgule et pipe pour les champs de type tableau (tags, artists, sponsors)."
    )

    doc.add_heading("5.3 Tarifs", level=2)
    doc.add_paragraph(
        "La colonne tarifs contient des valeurs numériques (ex : \u00ab 0.0 \u00bb "
        "pour gratuit). Il peut être utile de transformer \u00ab 0.0 \u00bb en "
        "\u00ab Gratuit \u00bb pour une meilleure lisibilité dans GEDEON."
    )

    doc.add_heading("5.4 Ville (non présente)", level=2)
    doc.add_paragraph(
        "Le CSV ne contient pas de colonne \u00ab ville \u00bb séparée. "
        "La ville est souvent incluse dans lieu_adresse "
        "(ex : \u00ab Place de la Nouvelle Aventure, Lille \u00bb) ou dans lieu_nom. "
        "Le géocodage automatique de GEDEON (basé sur l'adresse et les coordonnées GPS) "
        "devrait reconstituer la ville."
    )

    doc.add_page_break()

    # ==========================================================
    # 6. PROCÉDURE D'IMPORT
    # ==========================================================
    doc.add_heading("6. Procédure d'import pas à pas", level=1)

    doc.add_heading("6.1 Étape 1 : Sélection du fichier", level=2)
    for item in [
        "Se connecter à GEDEON Admin (https://gedeon-admin.onrender.com)",
        "Aller dans la section \u00ab Import \u00bb",
        "Glisser-déposer le fichier infolocale_20260315_8.csv dans la zone de dépôt",
        "Ou cliquer sur la zone pour sélectionner le fichier",
        "PapaParse détecte automatiquement le séparateur (virgule)",
        "Sélectionner l'utilisateur cible dans le menu déroulant",
        "Cliquer sur \u00ab Suivant \u00bb",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.2 Étape 2 : Mapping des champs", level=2)
    doc.add_paragraph(
        "Deux options sont possibles :"
    )

    p = doc.add_paragraph()
    run = p.add_run("Option A \u2013 Mapping IA (recommandé) : ")
    run.bold = True
    p.add_run(
        "Cliquer sur le bouton \u00ab Mapping IA \u00bb. L'IA analyse les noms de "
        "colonnes et les premières lignes de données pour proposer un mapping "
        "automatique. Vérifier les correspondances proposées et ajuster si nécessaire."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Option B \u2013 Mapping manuel : ")
    run.bold = True
    p.add_run(
        "Pour chaque colonne du CSV, sélectionner le champ GEDEON correspondant "
        "dans le menu déroulant. Se référer au tableau de correspondance de la "
        "section 4 de ce rapport."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Les 3 champs obligatoires doivent être mappés pour pouvoir continuer :"
    )
    for item in [
        "titre \u2192 title",
        "date_debut \u2192 begin_date",
        "lieu_nom \u2192 location_name",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.3 Étape 3 : Lancement de l'import", level=2)
    for item in [
        "Vérifier le résumé : nombre de lignes, type de mapping, utilisateur cible",
        "Cocher \u00ab Ignorer les lignes avec erreurs \u00bb (recommandé pour un premier import)",
        "Cocher \u00ab Tous les événements en privé \u00bb si souhaité",
        "Cliquer sur \u00ab Lancer l'import \u00bb",
        "Suivre la progression en temps réel (barre de progression, compteurs)",
        "L'import de 5 000 événements prend généralement quelques minutes",
        "À la fin, vérifier les statistiques : importés, doublons, erreurs, géocodés",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ==========================================================
    # 7. MAPPING IA
    # ==========================================================
    doc.add_heading("7. Utilisation du mapping IA", level=1)

    doc.add_paragraph(
        "GEDEON Admin propose un mapping assisté par IA qui analyse les noms "
        "de colonnes et les premières valeurs pour proposer automatiquement "
        "les correspondances."
    )

    doc.add_heading("7.1 Fonctionnement", level=2)
    for item in [
        "L'IA reçoit les noms de colonnes et les 5 premières lignes du CSV",
        "Elle reconnaît les patterns (titre \u2192 title, lieu_nom \u2192 location_name...)",
        "Elle détecte les formats composites (date+heure) et propose des transformations",
        "Le mapping est appliqué automatiquement dans l'interface",
        "Un bandeau explicatif décrit les choix de l'IA",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7.2 Mapping IA attendu pour ce fichier", level=2)
    doc.add_paragraph(
        "Avec le fichier Infolocale, le mapping IA devrait proposer :"
    )

    add_styled_table(
        doc,
        ["Colonne CSV", "Champ GEDEON", "Confiance"],
        [
            ["titre", "title", "Très élevée"],
            ["texte", "description", "Très élevée"],
            ["categorie", "category", "Élevée"],
            ["date_debut", "begin_date + start_time", "Élevée (split date/heure)"],
            ["date_fin", "end_date + end_time", "Élevée (split date/heure)"],
            ["lieu_nom", "location_name", "Très élevée"],
            ["lieu_adresse", "address", "Très élevée"],
            ["latitude", "latitude", "Très élevée"],
            ["longitude", "longitude", "Très élevée"],
            ["url", "website", "Élevée"],
            ["tarifs", "pricing", "Moyenne"],
        ],
        col_widths=[3.5, 5.5, 3],
    )

    doc.add_page_break()

    # ==========================================================
    # 8. CHAMPS NON IMPORTABLES
    # ==========================================================
    doc.add_heading("8. Champs non importables", level=1)

    doc.add_paragraph(
        "Certaines colonnes du CSV Infolocale n'ont pas d'équivalent dans "
        "le modèle de données GEDEON Admin et seront ignorées lors de l'import :"
    )

    add_styled_table(
        doc,
        ["Colonne CSV", "Raison de l'exclusion", "Action"],
        [
            ["objectID", "Identifiant Algolia interne", "Ignoré"],
            ["genre", "Pas de champ équivalent (sous-catégorie)", "Ignoré (categorie suffit)"],
            ["rubrique_lvl0", "Redondant avec categorie", "Ignoré"],
            ["rubrique_lvl1", "Hiérarchie de catégorie non supportée", "Ignoré"],
            ["photo_url", "Pas de champ image dans l'import", "Ignoré"],
            ["annule", "Pas de champ annulation", "Filtrer en pré-traitement"],
            ["complet", "Pas de champ complet", "Ignoré"],
            ["permanent", "Pas de champ permanent", "Ignoré"],
            ["source", "Métadonnée interne Infolocale", "Ignoré"],
            ["provider", "Métadonnée interne Infolocale", "Ignoré"],
            ["accessibilites", "Peut être mappé vers tags", "Optionnel"],
        ],
        col_widths=[3, 6, 4],
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Recommandation : ")
    run.bold = True
    p.add_run(
        "Avant l'import, filtrer les événements annulés (annule=True) "
        "pour ne pas importer des événements qui n'auront pas lieu."
    )

    doc.add_page_break()

    # ==========================================================
    # 9. POINTS D'ATTENTION
    # ==========================================================
    doc.add_heading("9. Points d'attention", level=1)

    doc.add_heading("9.1 Volume de données", level=2)
    doc.add_paragraph(
        "Le fichier contient 5 000 événements. L'import utilise le streaming SSE "
        "(Server-Sent Events) avec une progression ligne par ligne. "
        "Le timeout est fixé à 10 minutes, ce qui devrait suffire pour ce volume. "
        "En cas de coupure, les événements déjà importés restent en base."
    )

    doc.add_heading("9.2 Doublons", level=2)
    doc.add_paragraph(
        "GEDEON Admin détecte les doublons côté backend. Si un événement "
        "avec le même titre, la même date et le même lieu existe déjà, "
        "il sera marqué comme doublon et ne sera pas réimporté. "
        "Cela permet de relancer l'import en toute sécurité."
    )

    doc.add_heading("9.3 Géocodage automatique", level=2)
    doc.add_paragraph(
        "Le backend GEDEON effectue un géocodage automatique des événements "
        "importés. Comme le CSV Infolocale contient déjà les coordonnées GPS "
        "(latitude et longitude), le géocodage sera utilisé principalement "
        "pour enrichir les données (ville, département, pays)."
    )

    doc.add_heading("9.4 Pré-traitement optionnel du CSV", level=2)
    doc.add_paragraph(
        "Pour un import optimal, il est possible de pré-traiter le CSV avec "
        "un script Python avant l'import :"
    )
    for item in [
        "Supprimer les événements annulés (annule=True)",
        "Séparer date_debut en deux colonnes : date_debut_date et date_debut_heure",
        "Transformer tarifs \u00ab 0.0 \u00bb en \u00ab Gratuit \u00bb",
        "Extraire la ville depuis lieu_adresse (dernier élément après la virgule)",
        "Fusionner ages et accessibilites dans une colonne tags",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Cependant, le mapping IA de GEDEON Admin gère la plupart de ces "
        "transformations automatiquement. Le pré-traitement n'est donc pas "
        "obligatoire."
    )

    doc.add_heading("9.5 Import en mode privé", level=2)
    doc.add_paragraph(
        "L'option \u00ab Tous les événements en privé \u00bb permet d'importer "
        "les événements sans les rendre visibles immédiatement. C'est recommandé "
        "pour un premier import afin de vérifier la qualité des données avant "
        "de les publier."
    )

    doc.add_page_break()

    # ==========================================================
    # CONCLUSION
    # ==========================================================
    doc.add_heading("Conclusion", level=1)

    doc.add_paragraph(
        "L'import des 5 000 événements Infolocale dans GEDEON Admin est "
        "une opération simple grâce à l'interface d'import en 3 étapes "
        "et au mapping IA intégré.\n\n"
        "Les colonnes du CSV Infolocale correspondent bien aux champs attendus "
        "par GEDEON : les 3 champs obligatoires (titre, date_debut, lieu_nom) "
        "sont présents, ainsi que de nombreux champs optionnels (description, "
        "catégorie, coordonnées GPS, tarifs, URL).\n\n"
        "Le mapping IA devrait reconnaître automatiquement la majorité des "
        "correspondances. Seuls quelques champs spécifiques à Infolocale "
        "(objectID, annule, complet, source, provider) seront ignorés.\n\n"
        "Pour un import optimal, il est recommandé d'utiliser le mapping IA, "
        "de cocher \u00ab Ignorer les lignes avec erreurs \u00bb, et d'activer "
        "le mode privé pour vérifier les données avant publication."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- Fin du rapport ---")
    run.font.color.rgb = RGBColor(127, 127, 127)
    run.italic = True

    # ==========================================================
    # SAUVEGARDER
    # ==========================================================
    output_path = "doc/guide_import_infolocale_gedeon.docx"
    from pathlib import Path
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = generate_report()
    print(f"Rapport généré : {path}")
