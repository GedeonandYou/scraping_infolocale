"""Génère un rapport Word comparant les 3 approches d'accès aux événements Infolocale."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date


def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule."""
    from lxml import etree
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")


def add_styled_table(doc, headers, rows, col_widths=None):
    """Crée un tableau stylisé."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E74B5")

    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generate_report():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ==========================================================
    # PAGE DE TITRE
    # ==========================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT COMPARATIF")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Trois approches pour la récupération\n"
        "des événements Infolocale"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Projet : scraping_infolocale\n"
        f"Date : {date.today().strftime('%d/%m/%Y')}\n"
        f"Repository : github.com/GedeonandYou/scraping_infolocale"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()

    # ==========================================================
    # TABLE DES MATIÈRES
    # ==========================================================
    doc.add_heading("Table des matières", level=1)
    toc_items = [
        "1. Contexte et objectifs",
        "2. V1 \u2013 Scraping HTML (architecture originale)",
        "3. V2 \u2013 Scraping via API Algolia",
        "4. V3 \u2013 API à la demande (architecture proposée)",
        "5. Comparaison détaillée des 3 approches",
        "6. Analyse des risques",
        "7. Performance et scalabilité",
        "8. Coût et maintenance",
        "9. Recommandation finale",
        "10. Plan de migration",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ==========================================================
    # 1. CONTEXTE ET OBJECTIFS
    # ==========================================================
    doc.add_heading("1. Contexte et objectifs", level=1)

    doc.add_paragraph(
        "Le projet scraping_infolocale a pour objectif de récupérer les événements "
        "publiés sur la plateforme infolocale.fr (propriété du groupe Ouest-France) "
        "et de les mettre à disposition via une API REST et une interface frontend."
    )

    doc.add_heading("1.1 Source de données", level=2)
    doc.add_paragraph(
        "Infolocale.fr est une plateforme d'événements locaux couvrant l'ensemble "
        "du territoire français. Le site affiche les événements sous forme de cartes "
        "HTML et utilise en interne un index Algolia (memo_events, 110 000+ événements) "
        "pour la recherche côté frontend."
    )

    doc.add_heading("1.2 Les trois approches", level=2)
    doc.add_paragraph(
        "Ce rapport compare trois architectures possibles pour exploiter ces données :"
    )
    approaches = [
        (
            "V1 \u2013 Scraping HTML (originale)",
            "Parcourir les pages HTML d'infolocale.fr avec Selenium et "
            "BeautifulSoup pour extraire les événements, les stocker en base "
            "PostgreSQL, puis les servir via une API FastAPI.",
        ),
        (
            "V2 \u2013 Scraping Algolia",
            "Interroger directement l'API Algolia (index memo_events) pour "
            "récupérer massivement les événements, les stocker en base et les servir "
            "via l'API. Nécessite une clé API avec auto-renouvellement.",
        ),
        (
            "V3 \u2013 API à la demande (proposée)",
            "L'API FastAPI interroge Algolia en temps réel à chaque requête "
            "utilisateur, en utilisant les filtres géographiques natifs d'Algolia. "
            "Pas de stockage local des événements.",
        ),
    ]
    for label, text in approaches:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} : ")
        run.bold = True
        p.add_run(text)

    doc.add_page_break()

    # ==========================================================
    # 2. V1 - SCRAPING HTML
    # ==========================================================
    doc.add_heading("2. V1 \u2013 Scraping HTML (architecture originale)", level=1)

    doc.add_heading("2.1 Principe", level=2)
    doc.add_paragraph(
        "L'architecture originale du projet repose sur le scraping classique des pages "
        "HTML du site infolocale.fr. Un navigateur Chrome piloté par Selenium charge "
        "chaque page de résultats, puis BeautifulSoup extrait les informations des "
        "cartes d'événements."
    )

    doc.add_heading("2.2 Pipeline de données", level=2)
    flow_steps = [
        "Selenium charge une page de résultats sur infolocale.fr",
        "BeautifulSoup parse le HTML et extrait les cartes (.memo-card)",
        "Chaque carte est transformée en objet événement (titre, date, lieu, image\u2026)",
        "Les dates en français sont parsées (ex : \u00ab sam. 15 mars \u00bb \u2192 2026-03-15)",
        "Les événements sont stockés en base PostgreSQL (table scanned_events)",
        "Déduplication via UPSERT sur le champ uid (infolocale_{data_id})",
        "Enrichissement optionnel : géocodage via OpenRouteService",
        "Exposition via API FastAPI (GET /api/v1/events)",
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("2.3 Stack technique", level=2)
    add_styled_table(
        doc,
        ["Composant", "Technologie", "Rôle"],
        [
            ["Scraping", "Selenium 4 + BeautifulSoup", "Navigation et parsing HTML"],
            ["Driver", "webdriver-manager", "Gestion automatique de ChromeDriver"],
            ["API", "FastAPI + Uvicorn", "Serveur REST"],
            ["ORM", "SQLModel (SQLAlchemy + Pydantic)", "Modèles et validation"],
            ["Base de données", "PostgreSQL 15+", "Stockage événements"],
            ["Géocodage", "OpenRouteService", "Enrichissement géographique"],
            ["Migrations", "Alembic", "Schéma de base"],
            ["CLI", "Typer + Rich", "Interface en ligne de commande"],
        ],
        col_widths=[3.5, 5, 5],
    )

    doc.add_heading("2.4 Limites de cette approche", level=2)
    for item in [
        "Lent : ~50 événements par page, pagination manuelle, 2 s de délai entre chaque page",
        "Fragile : dépend de la structure HTML (sélecteurs CSS .memo-card, .gender, .day\u2026)",
        "WAF : le site est protégé par un WAF Ouest-France qui bloque les navigateurs headless",
        "Incomplet : seules les informations visibles sur la carte sont récupérées (pas de description complète)",
        "Volume limité : scraper 110k événements prendrait des heures",
        "Maintenance : tout changement de design du site casse les sélecteurs",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ==========================================================
    # 3. V2 - SCRAPING ALGOLIA
    # ==========================================================
    doc.add_heading("3. V2 \u2013 Scraping via API Algolia", level=1)

    doc.add_heading("3.1 Principe", level=2)
    doc.add_paragraph(
        "Cette approche contourne le HTML en interrogeant directement l'API Algolia "
        "utilisée par le frontend d'infolocale.fr. L'index memo_events contient "
        "l'intégralité des événements avec toutes leurs métadonnées (titre, texte, "
        "lieu, coordonnées GPS, dates, catégories, photos\u2026)."
    )

    doc.add_heading("3.2 Pipeline de données", level=2)
    flow_steps_v2 = [
        "Appel à l'endpoint Browse d'Algolia (pagination par cursor, sans limite de 1000)",
        "Récupération de tous les événements avec leurs métadonnées complètes",
        "Aplatissement des champs imbriqués (lieu, rubrique, _geoloc, photo)",
        "Stockage en base PostgreSQL ou export direct en CSV/JSON",
        "Exposition via API FastAPI",
    ]
    for i, step in enumerate(flow_steps_v2, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("3.3 Gestion de la clé API", level=2)
    doc.add_paragraph(
        "L'accès à l'API Algolia nécessite une \u00ab secured key \u00bb (base64) avec un "
        "champ validUntil qui expire régulièrement. Le service algolia_key_service.py "
        "gère le renouvellement automatique :"
    )
    for item in [
        "Lancement de Chrome headless avec options anti-détection (stealth)",
        "Chargement de infolocale.fr/activités pour déclencher les requêtes Algolia",
        "Interception des requêtes réseau via Chrome DevTools Protocol (CDP)",
        "Extraction de la clé depuis les paramètres d'URL (x-algolia-api-key)",
        "Gestion du WAF Ouest-France (détection de la page challenge, attente, retries)",
        "Renouvellement transparent : si la clé expire en cours d'export, elle est renouvelée automatiquement",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.4 Avantages par rapport à V1", level=2)
    for item in [
        "Beaucoup plus rapide : 1000 événements par requête vs ~50 par page HTML",
        "Données complètes : texte intégral, coordonnées GPS, tarifs, accessibilité\u2026",
        "Plus stable : pas de dépendance aux sélecteurs CSS du frontend",
        "Volume : peut récupérer les 110k+ événements en quelques minutes",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.5 Limites", level=2)
    for item in [
        "Clé API expirant régulièrement (nécessite Selenium pour le renouvellement)",
        "Dépendance à l'index Algolia d'Infolocale (s'il change, tout casse)",
        "Données pas forcément à jour entre deux syncs",
        "Nécessite toujours une base PostgreSQL pour le stockage",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ==========================================================
    # 4. V3 - API À LA DEMANDE
    # ==========================================================
    doc.add_heading("4. V3 \u2013 API à la demande (architecture proposée)", level=1)

    doc.add_heading("4.1 Principe", level=2)
    doc.add_paragraph(
        "Au lieu de scraper et stocker localement les 110k+ événements, "
        "l'API FastAPI interroge directement Algolia à chaque requête utilisateur. "
        "Les filtres géographiques natifs d'Algolia (aroundLatLng, aroundRadius) "
        "permettent de ne récupérer que les événements pertinents."
    )

    doc.add_heading("4.2 Flux de données", level=2)
    flow_steps_v3 = [
        "L'utilisateur envoie une requête : GET /events?lat=48.8&lng=2.3&radius=10km",
        "L'API FastAPI transforme la requête en appel Algolia avec les filtres natifs",
        "Algolia retourne uniquement les événements dans le rayon demandé",
        "L'API formate et retourne la réponse au frontend",
        "Cache optionnel : les résultats sont mis en cache (Redis ou in-memory)",
    ]
    for i, step in enumerate(flow_steps_v3, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4.3 Filtres Algolia disponibles", level=2)
    add_styled_table(
        doc,
        ["Paramètre Algolia", "Description", "Exemple"],
        [
            ["aroundLatLng", "Centre de recherche (lat, lng)", "48.8566,2.3522"],
            ["aroundRadius", "Rayon en mètres", "10000 (10 km)"],
            ["numericFilters", "Filtres sur dates", "date-first>=1710460800"],
            ["facetFilters", "Filtres sur catégories", "rubrique.lvl0:Concert"],
            ["hitsPerPage", "Nombre de résultats par page", "20"],
            ["page", "Numéro de page", "0"],
        ],
        col_widths=[4, 5, 4.5],
    )

    doc.add_paragraph()

    doc.add_heading("4.4 Composants supprimés par rapport à V1/V2", level=2)
    for item in [
        "scraper_service.py (scraping HTML Selenium \u2014 V1)",
        "storage_service.py (stockage PostgreSQL des événements)",
        "opendata_import_service.py (import CSV data.gouv.fr)",
        "export_events.py (export CSV/JSON)",
        "Table scanned_events (plus nécessaire pour les événements)",
        "Migrations Alembic pour la table événements",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.5 Composants conservés", level=2)
    for item in [
        "algolia_service.py (adapté pour les requêtes à la demande avec filtres géo)",
        "algolia_key_service.py (renouvellement automatique de clé \u2014 identique à V2)",
        "geocoding_service.py (géocodage inverse si nécessaire)",
        "API FastAPI (routes adaptées pour le proxy Algolia)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ==========================================================
    # 5. COMPARAISON DÉTAILLÉE
    # ==========================================================
    doc.add_heading("5. Comparaison détaillée des 3 approches", level=1)

    doc.add_heading("5.1 Tableau comparatif", level=2)
    add_styled_table(
        doc,
        ["Critère", "V1 \u2013 Scraping HTML", "V2 \u2013 Scraping Algolia", "V3 \u2013 API à la demande"],
        [
            [
                "Fraîcheur des données",
                "Obsolètes entre 2 syncs",
                "Obsolètes entre 2 syncs",
                "Temps réel",
            ],
            [
                "Richesse des données",
                "Limitée (carte HTML)",
                "Complète (toutes les métadonnées)",
                "Complète (toutes les métadonnées)",
            ],
            [
                "Volume stocké",
                "Variable (scraping lent)",
                "110k+ events en base",
                "Aucun (ou cache temporaire)",
            ],
            [
                "Latence utilisateur",
                "Rapide (SQL local)",
                "Rapide (SQL local)",
                "Moyenne (~200-500 ms)",
            ],
            [
                "Complexité infra",
                "PostgreSQL + Selenium + Cron",
                "PostgreSQL + Selenium (clé) + Cron",
                "FastAPI + Redis (optionnel)",
            ],
            [
                "Résilience WAF",
                "Bloquant (chaque page)",
                "Uniquement pour la clé API",
                "Uniquement pour la clé API",
            ],
            [
                "Filtrage géographique",
                "Post-traitement serveur",
                "Post-traitement serveur",
                "Natif Algolia (aroundLatLng)",
            ],
            [
                "Gestion de clé Algolia",
                "Non nécessaire",
                "Requise (auto-refresh)",
                "Requise (auto-refresh)",
            ],
            [
                "Mode offline",
                "Oui",
                "Oui",
                "Non",
            ],
            [
                "Maintenance",
                "Élevée (sélecteurs CSS)",
                "Moyenne (clé API)",
                "Faible",
            ],
            [
                "Conformité légale",
                "Zone grise (scraping)",
                "Zone grise (scraping massif)",
                "Plus léger (usage normal)",
            ],
        ],
        col_widths=[3, 3.5, 3.5, 3.5],
    )

    doc.add_paragraph()

    doc.add_heading("5.2 Fraîcheur des données", level=2)
    doc.add_paragraph(
        "Les approches V1 et V2 nécessitent une synchronisation régulière (cron job) "
        "pour maintenir les données à jour. Entre deux syncs, les événements "
        "ajoutés, modifiés ou annulés sur Infolocale ne sont pas reflétés.\n\n"
        "L'approche V3 retourne toujours les données les plus récentes puisqu'elle "
        "interroge Algolia en temps réel à chaque requête utilisateur."
    )

    doc.add_heading("5.3 Filtrage géographique", level=2)
    doc.add_paragraph(
        "Avec V1 et V2, le filtrage géographique est fait côté serveur après "
        "avoir récupéré tous les événements. Cela implique de stocker les "
        "coordonnées GPS et de calculer les distances.\n\n"
        "V3 utilise les filtres natifs d'Algolia (aroundLatLng, aroundRadius) "
        "qui effectuent ce calcul côté Algolia avec des performances optimisées "
        "(index géographique). Le frontend n'a qu'à envoyer la position de l'utilisateur."
    )

    doc.add_heading("5.4 Robustesse face au WAF", level=2)
    doc.add_paragraph(
        "V1 est la plus vulnérable : chaque page scrapée passe par le WAF "
        "Ouest-France, qui peut bloquer le navigateur headless à tout moment.\n\n"
        "V2 et V3 ne sont exposées au WAF que lors du renouvellement de la clé API "
        "(une seule visite de page nécessaire), ce qui réduit considérablement "
        "le risque de blocage."
    )

    doc.add_page_break()

    # ==========================================================
    # 6. ANALYSE DES RISQUES
    # ==========================================================
    doc.add_heading("6. Analyse des risques", level=1)

    add_styled_table(
        doc,
        ["Risque", "V1 \u2013 HTML", "V2 \u2013 Algolia", "V3 \u2013 API demande"],
        [
            [
                "Blocage WAF",
                "Critique\n(chaque page)",
                "Faible\n(clé uniquement)",
                "Faible\n(clé uniquement)",
            ],
            [
                "Changement design HTML",
                "Critique\n(sélecteurs cassent)",
                "Aucun impact",
                "Aucun impact",
            ],
            [
                "Expiration clé Algolia",
                "Aucun impact",
                "Moyen\n(auto-refresh géré)",
                "Moyen\n(auto-refresh géré)",
            ],
            [
                "Changement index Algolia",
                "Aucun impact",
                "Critique",
                "Critique",
            ],
            [
                "Rate limiting",
                "Élevé\n(nombreuses pages)",
                "Moyen\n(browse massif)",
                "Faible\n(quelques req/min)",
            ],
            [
                "Panne PostgreSQL",
                "Critique\n(plus de données)",
                "Critique\n(plus de données)",
                "Aucun impact\n(pas de base requise)",
            ],
            [
                "Conformité légale",
                "Risque élevé\n(scraping HTML)",
                "Risque moyen\n(scraping API)",
                "Risque faible\n(usage proxy)",
            ],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5],
    )

    doc.add_page_break()

    # ==========================================================
    # 7. PERFORMANCE ET SCALABILITÉ
    # ==========================================================
    doc.add_heading("7. Performance et scalabilité", level=1)

    doc.add_heading("7.1 Latence", level=2)
    add_styled_table(
        doc,
        ["Opération", "V1 \u2013 HTML", "V2 \u2013 Algolia", "V3 \u2013 API demande"],
        [
            [
                "Scraping initial",
                "Plusieurs heures\n(110k events)",
                "~5-15 min\n(110k events)",
                "Aucun",
            ],
            [
                "Requête utilisateur",
                "~10-50 ms\n(SQL local)",
                "~10-50 ms\n(SQL local)",
                "~200-500 ms\n(Algolia distant)",
            ],
            [
                "Recherche géographique",
                "~50-200 ms\n(calcul serveur)",
                "~50-200 ms\n(calcul serveur)",
                "~50-100 ms\n(index Algolia natif)",
            ],
            [
                "Renouvellement clé",
                "Non applicable",
                "~10-15 s\n(Selenium)",
                "~10-15 s\n(Selenium)",
            ],
        ],
        col_widths=[4, 3.5, 3.5, 3.5],
    )

    doc.add_paragraph()

    doc.add_heading("7.2 Scalabilité", level=2)
    doc.add_paragraph(
        "V1 : la scalabilité est très limitée. Le temps de scraping croît "
        "linéairement avec le nombre de pages, et le WAF peut bloquer à tout moment.\n\n"
        "V2 : meilleure scalabilité grâce à l'API Algolia, mais la taille de la base "
        "PostgreSQL et le temps de synchronisation restent des facteurs limitants.\n\n"
        "V3 : la scalabilité est gérée par Algolia (infrastructure CDN mondiale). "
        "L'ajout d'un cache Redis permet de réduire les appels répétés."
    )

    doc.add_heading("7.3 Cache (approche V3)", level=2)
    doc.add_paragraph(
        "L'approche V3 peut utiliser un cache à plusieurs niveaux :"
    )
    for item in [
        "Cache applicatif (in-memory, TTL 5 min) : pour les requêtes identiques rapprochées",
        "Cache Redis (TTL 30 min) : pour les requêtes par zone géographique",
        "Cache PostgreSQL (TTL 24 h, optionnel) : pour l'analytique et les stats",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ==========================================================
    # 8. COÛT ET MAINTENANCE
    # ==========================================================
    doc.add_heading("8. Coût et maintenance", level=1)

    doc.add_heading("8.1 Complexité du code", level=2)
    add_styled_table(
        doc,
        ["Composant", "V1 \u2013 HTML", "V2 \u2013 Algolia", "V3 \u2013 API demande"],
        [
            ["scraper_service.py", "~250 lignes", "Inutilisé", "Supprimé"],
            ["algolia_service.py", "Absent", "~300 lignes", "~150 (simplifié)"],
            ["algolia_key_service.py", "Absent", "~280 lignes", "~280 (inchangé)"],
            ["storage_service.py", "~150 lignes", "~150 lignes", "Supprimé"],
            ["opendata_import_service.py", "~200 lignes", "~200 lignes", "Supprimé"],
            ["export (CSV/JSON)", "~200 lignes", "~200 lignes", "Supprimé"],
            ["routes.py", "~200 lignes", "~200 lignes", "~100 (simplifié)"],
            ["models + migrations", "~150 lignes", "~150 lignes", "~50 (minimal)"],
            ["Total approximatif", "~1 150 lignes", "~1 480 lignes", "~580 lignes"],
        ],
        col_widths=[4, 3, 3, 3],
    )

    doc.add_paragraph()

    doc.add_heading("8.2 Infrastructure", level=2)
    add_styled_table(
        doc,
        ["Ressource", "V1 \u2013 HTML", "V2 \u2013 Algolia", "V3 \u2013 API demande"],
        [
            ["PostgreSQL", "Requis", "Requis", "Optionnel (cache)"],
            ["Redis", "Non utilisé", "Non utilisé", "Optionnel (cache)"],
            ["Chrome/Selenium", "Requis (scraping)", "Requis (clé API)", "Requis (clé API)"],
            ["Cron job", "Requis (sync)", "Requis (sync)", "Non requis"],
            ["Espace disque", "~500 Mo", "~500 Mo", "Minimal"],
        ],
        col_widths=[4, 3, 3.5, 3.5],
    )

    doc.add_page_break()

    # ==========================================================
    # 9. RECOMMANDATION FINALE
    # ==========================================================
    doc.add_heading("9. Recommandation finale", level=1)

    doc.add_heading("9.1 Synthèse", level=2)
    doc.add_paragraph(
        "V1 (scraping HTML) a permis de démarrer le projet et de valider le concept, "
        "mais cette approche est fragile, lente et limitée en données.\n\n"
        "V2 (scraping Algolia) a apporté un gain majeur en termes de vitesse, de "
        "complétude des données et de stabilité. Le mécanisme d'auto-renouvellement "
        "de la clé API résout le problème de l'expiration.\n\n"
        "V3 (API à la demande) est l'évolution logique pour un cas d'usage centré "
        "sur la recherche d'événements par localisation : elle est plus simple, "
        "garantit des données fraîches, et exploite le filtrage géographique natif d'Algolia."
    )

    doc.add_heading("9.2 Approche recommandée : Hybride V2+V3", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'approche recommandée combine les avantages de V2 et V3 :"
    )

    doc.add_paragraph()

    rec_items = [
        (
            "API à la demande pour les requêtes utilisateur (V3)",
            "Les recherches géographiques et par catégorie sont transmises "
            "directement à Algolia avec les filtres natifs (aroundLatLng, facetFilters). "
            "Données toujours fraîches, filtrage performant.",
        ),
        (
            "Cache intelligent",
            "Un cache Redis (TTL 15-30 min) stocke les résultats des requêtes "
            "fréquentes pour réduire la latence et éviter de surcharger Algolia.",
        ),
        (
            "Renouvellement automatique de la clé (V2)",
            "Le service algolia_key_service.py gère le renouvellement "
            "transparent de la clé API expirée, commun à V2 et V3.",
        ),
        (
            "Export massif optionnel (V2)",
            "Conserver la possibilité d'export CSV/JSON via algolia_service.py "
            "pour les besoins d'analytique ou d'archivage.",
        ),
    ]
    for label, text in rec_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} : ")
        run.bold = True
        run.font.color.rgb = RGBColor(46, 116, 181)
        p.add_run(text)

    doc.add_heading("9.3 Matrice de décision", level=2)
    add_styled_table(
        doc,
        ["Cas d'usage", "Approche recommandée"],
        [
            ["App web : chercher des événements proches", "V3 \u2013 API à la demande"],
            ["App mobile avec mode offline", "V2 \u2013 Scraping Algolia + sync"],
            ["Dashboard analytique", "V2 \u2013 Scraping Algolia (données locales)"],
            ["Moteur de recommandation", "Hybride V2+V3 (cache + ML local)"],
            ["Simple affichage d'événements", "V3 \u2013 API à la demande"],
            ["Export massif CSV/JSON", "V2 \u2013 Scraping Algolia"],
            ["Prototype rapide", "V1 \u2013 Scraping HTML (déjà en place)"],
        ],
        col_widths=[7, 7],
    )

    doc.add_page_break()

    # ==========================================================
    # 10. PLAN DE MIGRATION
    # ==========================================================
    doc.add_heading("10. Plan de migration vers V3", level=1)

    doc.add_paragraph(
        "Si la décision est prise de migrer vers l'approche V3 (API à la demande), "
        "voici les étapes recommandées :"
    )

    phases = [
        (
            "Phase 1 : Adapter AlgoliaService (1-2 jours)",
            [
                "Ajouter les méthodes de recherche géographique (aroundLatLng, aroundRadius)",
                "Ajouter le support des facetFilters pour les catégories",
                "Transformer les hits Algolia en schéma EventRead existant",
            ],
        ),
        (
            "Phase 2 : Créer les nouveaux endpoints API (1-2 jours)",
            [
                "GET /events?lat=...&lng=...&radius=... (recherche géo)",
                "GET /events?category=...&city=... (filtres)",
                "Conserver la pagination (page, page_size)",
            ],
        ),
        (
            "Phase 3 : Ajouter le cache (1 jour)",
            [
                "Intégrer Redis comme cache de résultats",
                "Définir les clés de cache (hash de la requête)",
                "Configurer les TTL par type de requête",
            ],
        ),
        (
            "Phase 4 : Adapter le frontend (1-2 jours)",
            [
                "Mettre à jour les appels API pour utiliser les nouveaux endpoints",
                "Passer la géolocalisation du navigateur à l'API",
                "Gérer le chargement (indicateur pendant l'appel Algolia)",
            ],
        ),
        (
            "Phase 5 : Nettoyage (1 jour)",
            [
                "Supprimer scraper_service.py (V1 obsolète)",
                "Simplifier les migrations Alembic",
                "Mettre à jour la documentation",
                "Conserver algolia_service.py et export_events.py pour l'export ponctuel",
            ],
        ),
    ]

    for phase_title, steps in phases:
        doc.add_heading(phase_title, level=2)
        for step in steps:
            doc.add_paragraph(step, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Effort total estimé : 5 à 8 jours de développement")
    run.bold = True

    doc.add_page_break()

    # ==========================================================
    # CONCLUSION
    # ==========================================================
    doc.add_heading("Conclusion", level=1)

    doc.add_paragraph(
        "Le projet a évolué en trois étapes naturelles :\n\n"
        "V1 (scraping HTML) a permis de valider le concept et de construire "
        "l'infrastructure de base (API, base de données, modèles).\n\n"
        "V2 (scraping Algolia) a apporté un saut qualitatif majeur : accès à "
        "l'intégralité des 110k+ événements avec toutes leurs métadonnées, "
        "et un mécanisme robuste de renouvellement de clé API.\n\n"
        "V3 (API à la demande) est l'évolution naturelle pour un produit centré "
        "sur la recherche d'événements par localisation. Elle simplifie "
        "drastiquement l'architecture tout en garantissant des données toujours "
        "à jour et un filtrage géographique performant.\n\n"
        "L'approche hybride V2+V3 recommandée permet de conserver le meilleur de "
        "chaque version selon les besoins."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- Fin du rapport ---")
    run.font.color.rgb = RGBColor(127, 127, 127)
    run.italic = True

    # ==========================================================
    # SAUVEGARDER
    # ==========================================================
    output_path = "doc/rapport_comparatif_scraping_vs_api.docx"
    from pathlib import Path
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = generate_report()
    print(f"Rapport généré : {path}")
